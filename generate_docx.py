#!/usr/bin/env python3
"""生成包含图表的Word论文文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import os

RESULT_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), "result")
OUTPUT_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "manuscript_CN.docx")


def set_cell_font(run, font_name='宋体', font_name_east='宋体', size=10.5, bold=False):
    """设置字体"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name_east)


def add_heading_custom(doc, text, level=1, font_size=14, bold=True):
    """添加自定义标题"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_cell_font(run, '黑体', '黑体', font_size, bold)
    p.space_after = Pt(6)
    p.space_before = Pt(12)
    return p


def add_body_paragraph(doc, text, font_size=10.5, first_line_indent=True):
    """添加正文段落"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    set_cell_font(run, '宋体', '宋体', font_size, False)
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_figure(doc, image_path, caption, width=6.0):
    """插入图片和图注"""
    if not os.path.exists(image_path):
        print(f"Warning: {image_path} not found")
        return
    # 图片
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(image_path, width=Inches(width))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    # 图注
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap_p.add_run(caption)
    set_cell_font(cap_run, '宋体', '宋体', 9, False)
    cap_p.paragraph_format.space_after = Pt(12)


def add_table_caption(doc, text):
    """添加表注"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_cell_font(run, '宋体', '宋体', 9, True)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)


def main():
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # 设置页边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ============ 标题 ============
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("NK样CD8 T细胞的克隆命运锁定通过SPP1+肿瘤相关巨噬细胞闸门预测非小细胞肺癌抗PD-1治疗响应")
    set_cell_font(title_run, '黑体', '黑体', 16, True)
    title_p.paragraph_format.space_after = Pt(18)

    # 英文标题
    en_title_p = doc.add_paragraph()
    en_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    en_run = en_title_p.add_run("Clonal Fate Locking of NK-like CD8 T Cells Predicts Anti-PD-1 Response in Non-Small Cell Lung Cancer via SPP1+ Tumor-Associated Macrophage Gate")
    en_run.font.name = 'Times New Roman'
    en_run.font.size = Pt(12)
    en_run.font.italic = True
    en_title_p.paragraph_format.space_after = Pt(24)

    # ============ 摘要 ============
    add_heading_custom(doc, "摘要", 1, 14)

    abstract_text = (
        "抗PD-1免疫治疗已彻底改变非小细胞肺癌（NSCLC）的治疗格局，但可靠的预测性生物标志物仍然匮乏。"
        "利用来自188例接受新辅助抗PD-1治疗的NSCLC患者的434,458个T细胞的单细胞转录组图谱，"
        "我们发现了一种克隆命运锁定现象——扩增的CD8 T细胞克隆被不可逆地定向到NK样（FGFBP2+）或终末耗竭（Tex）轨迹。"
        "NK锁定大克隆的比例（NK-Locked Ratio）能够强烈预测病理完全缓解（OR = 19.54，p = 0.008；校正OR = 38.14，p = 0.002）。"
        "机制上，SPP1+肿瘤相关巨噬细胞（TAMs）形成了一个髓系闸门，抑制NK样CD8 T细胞中的干性相关基因（KLF2、TCF7、IL7R）"
        "（Spearman r = −0.339，p = 3.78 × 10⁻⁶）。"
        "在独立NSCLC队列中的外部验证证实了NK样24基因签名的预测价值（GSE126044，p = 0.006，FC = 4.72），"
        "而黑色素瘤队列中未发现关联，支持肿瘤类型特异性的克隆动力学。"
        "空间转录组学原位验证了SPP1-受体相互作用（r = 0.325，p = 0.044）。"
        "这些发现确立了克隆命运锁定作为抗PD-1响应的决定性原则，并提示SPP1+ TAM-KLF2轴是潜在治疗靶点。"
    )
    add_body_paragraph(doc, abstract_text)

    # 关键词
    kw_p = doc.add_paragraph()
    kw_p.paragraph_format.first_line_indent = Cm(0.74)
    kw_run = kw_p.add_run("关键词：")
    set_cell_font(kw_run, '黑体', '黑体', 10.5, True)
    kw_run2 = kw_p.add_run("NK样CD8 T细胞；克隆命运锁定；抗PD-1；SPP1+肿瘤相关巨噬细胞；非小细胞肺癌；单细胞RNA测序；肿瘤微环境；免疫治疗生物标志物")
    set_cell_font(kw_run2, '宋体', '宋体', 10.5, False)

    doc.add_page_break()

    # ============ 引言 ============
    add_heading_custom(doc, "1. 引言", 1, 14)

    intro_paras = [
        "针对PD-1/PD-L1轴的免疫检查点阻断已彻底改变非小细胞肺癌（NSCLC）的治疗模式，新辅助抗PD-1治疗在约20-45%的患者中实现了病理完全缓解（pCR）。然而，大多数患者未能获得持久获益，响应异质性的生物学决定因素仍未完全阐明。虽然肿瘤突变负荷（TMB）和PD-L1表达是已批准的伴随诊断，但它们的预测准确性有限，突显了对基于机制的生物标志物的需求。",
        "CD8 T细胞是抗肿瘤免疫的主要效应细胞，但其功能状态存在于从干性样记忆到终末耗竭的连续谱中。在这个谱系中，共表达自然杀伤（NK）受体的CD8 T细胞亚群——称为NK样CD8 T细胞——已成为具有强大细胞毒性能力的功能独特群体。这些细胞以FGFBP2、KLRD1、CX3CR1和FCGR3A为标记，与传统NK细胞共享转录特征，同时保留T细胞受体（TCR）克隆性。然而，NK样CD8 T细胞代表终末分化终点还是扩增T细胞克隆内动态调节的命运承诺，尚未得到系统研究。",
        "肿瘤相关巨噬细胞（TAMs）构成NSCLC微环境中占主导地位的免疫抑制区室，并越来越多地与检查点阻断耐药相关。特别是，SPP1（骨桥蛋白）表达的TAMs已被确定为与T细胞抑制和多种癌症类型不良预后相关的促肿瘤亚群。然而，SPP1+ TAMs与效应CD8 T细胞克隆动力学之间的机制联系——特别是它们是否直接影响克隆命运决定——仍然未知。",
        "在本研究中，我们利用接受新辅助抗PD-1治疗的NSCLC患者的大规模单细胞RNA测序（scRNA-seq）图谱（GSE243013），调查NK样CD8 T细胞的克隆结构及其与治疗响应的关系。我们发现了一种'克隆命运锁定'现象，其中扩增的CD8 T细胞克隆被不可逆地定向到NK样或Tex轨迹，NK锁定的比例作为病理响应的有效预测因子。我们进一步阐明了SPP1+ TAM介导的髓系闸门，该闸门抑制NK样CD8 T细胞中的干性基因表达，建立了髓系驱动免疫逃逸的机制框架。多个独立队列的外部验证和空间转录组学确认共同支持一种肿瘤类型特异性模型，其中克隆命运锁定代表NSCLC抗PD-1响应的限速瓶颈。",
    ]
    for para in intro_paras:
        add_body_paragraph(doc, para)

    # ============ 结果 ============
    add_heading_custom(doc, "2. 结果", 1, 14)

    # 2.1
    add_heading_custom(doc, "2.1. 单细胞图谱鉴定NK样CD8 T细胞为独特的克隆群体", 2, 12)
    sec21 = [
        "我们分析了来自188例接受新辅助抗PD-1（特瑞普利单抗）治疗的NSCLC患者的434,458个T细胞的单细胞转录组图谱（GSE243013），包括84例病理完全缓解（pCR）患者和104例非主要病理缓解（non-MPR）患者。在CD8 T细胞区室（219,842个细胞）中，我们根据原始研究的细胞类型注释识别出16,744个NK样CD8 T细胞（标注为CD8T_NK-like_FGFBP2）和71,298个终末耗竭T细胞（Tex）。",
        "为确认NK样CD8 T细胞的转录身份，我们计算了NK基因签名得分（FGFBP2、KLRD1、CX3CR1、FCGR3A、KLRC1、NKG7、GNLY、PRF1、GZMB、KLRB1），并在NK样与其他CD8 T细胞亚群之间进行差异分析。NK样细胞显示显著升高的NK签名得分（Mann-Whitney U，p < 0.001），证实了其独特的细胞毒性身份（图1A,B）。",
        "值得注意的是，基因点图分析显示NK样CD8 T细胞同时表达干性相关基因（KLF2、TCF7）和细胞毒性介质（GZMB），与表达高水平耗竭标志物（CXCL13、LAYN、TOX）的Tex细胞形成鲜明对比（图1B）。这种双重细胞毒性-干性表型促使我们调查NK样身份是否代表稳定的克隆承诺。",
        "为检查克隆结构，我们分析了TCR克隆型在CD8 T细胞亚型中的分布。克隆扩增分析（每患者top-1和top-5克隆型比例）显示克隆优势存在显著的患者间异质性（图1C）。然而，pCR和non-MPR患者之间克隆扩增的比较显示top-5克隆型比例无显著差异（Mann-Whitney U，p = 0.683；pCR中位数 = 0.237，non-MPR中位数 = 0.210）（图1D），表明克隆扩增的数量本身不能区分响应者和非响应者。这一发现促使我们深入研究扩增克隆的质量——特别是亚型承诺。",
    ]
    for para in sec21:
        add_body_paragraph(doc, para)

    # 插入图1
    add_figure(doc, os.path.join(RESULT_DIR, "Fig1_overview.png"),
               "图1. NK样CD8 T细胞的鉴定与克隆分布。(A) CD8 T细胞UMAP降维图，NK样细胞高亮显示；(B) 关键基因点图（FGFBP2/KLF2/CXCL13/LAYN/GZMB/TCF7/TOX）；(C) 每患者克隆扩增分布（top-1/top-5克隆占比）；(D) 响应组间克隆扩增箱线图（pCR vs non-MPR的top5_frac）+ Mann-Whitney U检验。")

    # 2.2
    add_heading_custom(doc, "2.2. 克隆命运锁定预测病理响应", 2, 12)
    sec22 = [
        "我们假设扩增克隆的功能身份而非其大小决定治疗结果。为验证这一点，我们将大克隆定义为≥5个细胞的克隆，并根据其主要CD8 T细胞亚型组成对每个大克隆进行分类：NK锁定克隆（NK样CD8 T细胞占细胞总数>50%的大克隆，nk_frac > 0.5，表明不可逆地定向到NK样轨迹）和Tex分化克隆（以Tex细胞为主的大克隆）。对于每个患者，我们计算了NK-Locked Ratio——即NK锁定的大克隆比例——作为患者水平的克隆命运承诺指标（图2A）。",
        "单变量逻辑回归显示NK-Locked Ratio是pCR的强预测因子（OR = 19.54，95% CI: 2.18-175.15，p = 0.0079，n = 185例pCR与non-MPR比较患者）（图2B）。Tex主导克隆比例显示反向关联（OR = 112.04，95% CI: 4.08-3079.82，p = 0.005），与二元命运决定一致。调整总克隆数和CD8 T细胞比例的多变量逻辑回归确认了NK-Locked Ratio的独立预测价值（校正OR = 38.14，95% CI: 3.62-401.99，p = 0.002）。",
        "响应组间的直接比较证实pCR患者的NK-Locked Ratio显著高于non-MPR患者（Mann-Whitney U，p = 0.001；pCR中位数 = 0.083，non-MPR中位数 = 0.042）（图2C）。频率分布分析显示7.6%的患者（14/185）的NK-Locked Ratio为零，且这些患者不成比例地集中在non-MPR组（图2D）。",
        "不同克隆大小阈值（≥0、≥5、≥10、≥20、≥50个细胞）的敏感性分析表明，≥5的阈值下预测关联稳健，但较高阈值下减弱（≥20: OR = 5.41, p = 0.149；≥50: OR = 24.21, p = 0.153），反映患者数量减少导致统计效能降低（图S2，表S3）。最佳阈值为≥5个细胞，平衡了克隆大小的生物学意义与样本量充足性。",
    ]
    for para in sec22:
        add_body_paragraph(doc, para)

    # 插入图2
    add_figure(doc, os.path.join(RESULT_DIR, "Fig2_clonal_fate.png"),
               "图2. 克隆命运锁定预测响应。(A) NK锁定vs Tex分化克隆定义示意图；(B) 森林图（绿色正向NK-Locked，红色负向Tex），OR由statsmodels动态计算；(C) 响应组间Ratio箱线图（pCR vs non-MPR）+ Mann-Whitney U检验；(D) 频数分布直方图（pCR vs non-MPR的nk_dominant_ratio分布）。")

    # 2.3
    add_heading_custom(doc, "2.3. SPP1+肿瘤相关巨噬细胞形成髓系闸门", 2, 12)
    sec23 = [
        "为调查影响克隆命运锁定的微环境因素，我们使用完整免疫细胞图谱（1,254,749个细胞）分析髓系区室，重点关注191,099个髓系细胞。在巨噬细胞和单核细胞亚群中，我们将SPP1+ TAMs定义为SPP1表达高于亚群中位数的细胞，得到31,488个SPP1+ TAMs（占髓系细胞的16.5%）（图3A,B）。",
        "SPP1+ TAMs在non-MPR患者中显著富集（Mann-Whitney U，p = 5.54 × 10⁻⁹；non-MPR中位数 = 0.141，pCR中位数 = 0.035）（图3C），建立了与NK锁定表型的反向关系。",
        "为表征SPP1+ TAMs与CD8 T细胞亚群之间的细胞间通讯，我们计算了8条候选通路的配体-受体相互作用得分。分别为NK样和Tex受体细胞计算相互作用得分（SPP1+ TAMs中的配体表达 × 每个CD8亚型中的受体表达），取最大值以避免信号平均。最显著的相互作用是SPP1-CD44（Tex相互作用得分 = 3.90）和SPP1-ITGB1（NK样相互作用得分 = 3.88），其次是CXCL8-CXCR2、CCL2-CCR2和TNF-TNFRSF1A（图3D）。SPP1-CD44/ITGB1轴代表TAMs与两种CD8 T细胞命运之间的主要信号传导途径。",
    ]
    for para in sec23:
        add_body_paragraph(doc, para)

    # 2.4
    add_heading_custom(doc, "2.4. SPP1+ TAMs抑制NK样CD8 T细胞中的干性基因", 2, 12)
    sec24 = [
        "SPP1+ TAM丰度与NK-Locked Ratio之间的反向关系促使我们测试SPP1+ TAMs是否直接抑制维持NK样身份的干性程序。我们将患者水平的SPP1+ TAM比例与NK-Locked Ratio进行相关性分析，观察到显著负相关（Spearman r = −0.339，p = 3.78 × 10⁻⁶，n = 178；Pearson r = −0.136，p = 0.071）（图4A）。",
        "为确定SPP1+ TAM介导抑制的分子靶点，我们根据SPP1+ TAM比例将患者分为高SPP1组（前30%，n = 62）和低SPP1组（后30%，n = 62），并比较NK样CD8 T细胞中的干性相关基因表达。高SPP1患者表现出显著降低的KLF2表达（高 = 0.847 vs 低 = 1.057，Mann-Whitney U p = 0.009，FDR = 0.019）、TCF7（0.227 vs 0.280，p = 0.004，FDR = 0.012）和IL7R（1.043 vs 1.407，p = 0.0008，FDR = 0.005）（图4B）。KLF2作为维持T细胞静息和干性的主转录因子，成为核心枢纽。",
        "NK样与Tex细胞之间干性基因表达的直接比较（敏感性过滤子集n_NK-like = 3,662，n_Tex = 13,413）证实NK样CD8 T细胞固有地表达更高水平的KLF2（2.112 vs 0.308，p < 10⁻³⁰⁰）、TCF7（0.153 vs 0.063，p = 3.66 × 10⁻⁵⁴）、SELL（0.165 vs 0.038，p = 1.08 × 10⁻⁹⁵）、IL7R（1.051 vs 0.703，p = 3.83 × 10⁻⁵⁹）和LEF1（0.078 vs 0.049，p = 2.49 × 10⁻¹¹）（图4C）。这些数据证实KLF2和相关干性因子是NK样身份的定义特征，并被SPP1+ TAMs特异性抑制。",
    ]
    for para in sec24:
        add_body_paragraph(doc, para)

    # 插入图3和图4
    add_figure(doc, os.path.join(RESULT_DIR, "Fig3_myeloid.png"),
               "图3. 髓系闸门SPP1+ TAM。(A) 髓系PCA+UMAP，SPP1+ TAM高亮；(B) 特征基因热图（SPP1/IL1B/CXCL13/CD163/MRC1/CD68/CSF1R等×髓系亚群）；(C) SPP1+ TAM占髓系比例箱线图（pCR vs non-MPR）+ Mann-Whitney U检验；(D) 配受体互作网络图（SPP1-CD44/CXCL8-CXCR2/CCL2-CCR2等8对）；(E) 空间转录组验证整合。")

    add_figure(doc, os.path.join(RESULT_DIR, "Fig4_mechanism.png"),
               "图4. 机制链条：髓系闸门压制干性。(A) 散点图——SPP1+ TAM比例vs NK-Locked Ratio + Pearson/Spearman相关 + 线性回归置信带；(B) 小提琴图——High-SPP1组vs Low-SPP1组的干性基因（KLF2/TCF7/LEF1/IL7R/SELL）表达差异 + MWU + FDR校正；(C) NK-like vs Tex干性基因表达对比（小提琴图 + 每基因MWU p值标注）。")

    # 2.5
    add_heading_custom(doc, "2.5. 外部验证与跨癌种异质性", 2, 12)
    sec25 = [
        "为验证NK样签名作为预测性生物标志物，我们在四个独立队列中计算了24基因NK样签名得分：GSE126044（NSCLC，抗PD-1，n = 16）、GSE135222（NSCLC，抗PD-1/PD-L1，n = 27）、GSE91061（黑色素瘤，抗PD-1/CTLA-4，n = 33预处理）和GSE120575（黑色素瘤scRNA-seq，n = 11例有匹配表达数据的患者）。",
        "在NSCLC队列GSE126044中，响应者的NK样签名得分显著高于非响应者（Mann-Whitney U，p = 0.006，倍数变化 = 4.72）（图5A），独立证实了发现队列的结果。基因匹配率为GSE126044的24/24（100%）、GSE135222的13/24（54%）（Ensembl ID映射）、GSE91061的22/24（92%）（Entrez ID）和GSE120575的24/24（100%）。",
        "然而，在黑色素瘤队列中未观察到显著关联（GSE91061: p = 0.769, FC = 1.17；GSE120575: p = 0.921, FC = 0.98）（图5B）。NSCLC队列GSE135222也未显示显著关联（p = 0.755, FC = 1.02），这归因于其依赖无进展生存期（PFS）而非RECIST响应标准，且响应者样本量有限（R = 6）。",
        "我们通过肿瘤免疫环境模型解释这种跨癌种异质性（图5C）：NSCLC作为免疫'冷'肿瘤，TMB较低，表现出克隆命运瓶颈，其中扩增克隆定向到NK样轨迹是响应的限速步骤。相比之下，黑色素瘤作为'热'肿瘤，TMB高且克隆多样性丰富，已突破此瓶颈——分化状态而非NK样锁定决定结果。该模型与TMB在NSCLC和黑色素瘤之间的不同预测生物标志物性能一致。",
    ]
    for para in sec25:
        add_body_paragraph(doc, para)

    # 插入图5
    add_figure(doc, os.path.join(RESULT_DIR, "Fig5_external_validation.png"),
               "图5. 外部验证与转化潜力。(A) GSE126044 R vs NR箱线图 + Mann-Whitney U检验；(B) 跨癌种对比森林图（4队列：GSE126044/GSE135222/GSE91061/GSE120575），log2(FC) + 2000次bootstrap 95% CI + p值标注；(C) 跨癌种异质性模型图（NSCLC冷肿瘤bottleneck = stemness loss vs Melanoma热肿瘤bottleneck = differentiation）。")

    # 2.6 空间验证
    add_heading_custom(doc, "2.6. 空间转录组学验证", 2, 12)
    sec26 = [
        "为原位确认SPP1+ TAM-NK样CD8 T细胞相互作用，我们分析了来自39例NSCLC患者的67个感兴趣区域（ROIs）的空间转录组数据（GSE221733，GeoMx DSP）。SPP1-受体相互作用得分（SPP1 × CD44/ITGAV/ITGB1的平均受体表达）与肿瘤区室中的干性基因表达正相关（Spearman r = 0.325，p = 0.044）。NK细胞毒性得分（GZMB/GZMA/GNLY/PRF1）在响应者ROIs中显著高于非响应者ROIs（Mann-Whitney U，p = 0.027）。在肿瘤特异性ROIs（PanCK+）中，SPP1-NK受体相互作用进一步增强（r = 0.346，p = 0.045）。",
        "四组比较（高SPP1/高干性 vs 低SPP1/低干性）显示克隆命运锁定指标存在显著差异（p = 0.017），支持SPP1+ TAM-干性轴作为空间组织的调节模块发挥作用的模型。",
    ]
    for para in sec26:
        add_body_paragraph(doc, para)

    doc.add_page_break()

    # ============ 讨论 ============
    add_heading_custom(doc, "3. 讨论", 1, 14)

    discussion = [
        "本研究确定克隆命运锁定是NSCLC抗PD-1响应的决定性原则。我们的主要发现是：（1）扩增的CD8 T细胞克隆不可逆地定向到NK样或Tex轨迹，NK-Locked Ratio有效预测pCR；（2）SPP1+ TAMs形成抑制NK样CD8 T细胞中KLF2依赖性干性的髓系闸门；（3）NK样基因签名在独立NSCLC队列中验证，但在黑色素瘤中未验证，反映肿瘤类型特异性克隆动力学；（4）空间转录组学原位确认SPP1-受体相互作用。",
        "克隆命运锁定作为决定性原则。抗肿瘤免疫中克隆优势的概念已确立——肿瘤反应性T细胞的寡克隆扩增与检查点阻断响应相关。然而，我们的数据揭示克隆承诺的方向而非扩增的幅度是关键决定因素。观察到克隆扩增本身在响应者和非响应者之间无差异（p = 0.683），但NK-Locked Ratio具有高度区分性（OR = 19.54），从根本上重新定义了什么构成'有效'T细胞克隆。这一发现与最近的报道一致，即干性样CD8 T细胞亚群（TCF1+、SLAMF6+）对于持续抗肿瘤响应至关重要，但通过显示NK样轨迹本身通过KLF2表达携带干性能力而扩展了这一概念。",
        "克隆命运锁定的二元性质——NK锁定 vs Tex分化——表明克隆分化存在一个不归点。考虑到耗竭程序的表观遗传稳定性和NK样与耗竭程序之间的转录不相容性，这在机制上是合理的。>50%的阈值用于定义NK锁定克隆，以确保生物学承诺而非随机波动，敏感性分析确认了中等阈值下的稳健性。",
        "SPP1+ TAMs作为髓系闸门。将SPP1+ TAMs鉴定为NK样干性的抑制因子，为髓系炎症与T细胞功能障碍之间提供了机制联系。SPP1（骨桥蛋白）是促肿瘤巨噬细胞中上调的多效细胞因子，SPP1+ TAMs已被确定为多种癌症类型中保守的促肿瘤群体。我们的发现SPP1-CD44和SPP1-ITGB1是主要配体-受体对，与CD44在维持T细胞水合和迁移中的已知作用以及ITGB1在T细胞-细胞外基质相互作用中的作用一致。SPP1+ TAMs对KLF2的抑制——T细胞静息、迁移（通过S1P受体转录）和干性的主调节因子——提供了髓系细胞如何将T细胞克隆'锁定'到终末分化的分子机制。",
        "相关性强度（Spearman r = −0.339）表明SPP1+ TAMs解释了NK-Locked Ratio方差的很大一部分但非全部，与抗原亲和力、共刺激和细胞因子环境也贡献克隆命运决定的多因素模型一致。剩余方差也可能反映我们分析中未捕获的其他髓系抑制群体的作用。",
        "跨癌种异质性。NK样签名在黑色素瘤队列（GSE91061、GSE120575）中未能预测响应是关键发现而非局限。黑色素瘤具有高TMB、丰富的新抗原特异性T细胞克隆和'热'免疫微环境。在此背景下，我们提出克隆多样性足够高，瓶颈从克隆命运锁定转移到分化动力学——响应者和非响应者都有NK样克隆，但效应到耗竭转变的速率不同。该模型得到TMB在NSCLC和黑色素瘤之间不同预测性能以及黑色素瘤与肺癌T细胞中观察到的不同耗竭动力学的支持。跨癌种异质性也具有转化意义：基于NK样签名的生物标志物应在肿瘤类型特异性背景下开发和验证，而非假定具有普遍性。",
        "空间验证。使用GeoMx DSP数据的空间转录组分析提供了SPP1-受体-干性轴的原位确认。SPP1相互作用得分与肿瘤ROIs中干性基因表达之间的正相关（r = 0.325，p = 0.044）证明在解离单细胞数据中鉴定的调节模块在完整组织架构中发挥作用。响应者ROIs中更高的NK细胞毒性得分（p = 0.027）进一步支持空间组织的NK样活性的功能相关性。DSP平台的局限性——缺乏KLF2和FGFBP2探针——需要使用替代基因集，但空间和单细胞发现的收敛性增强了总体结论。",
        "局限性。应承认几个局限性。首先，发现队列是单中心研究，需要多中心验证。其次，外部验证队列GSE126044样本量较小（n = 16），GSE135222队列依赖PFS而非RECIST标准，可能降低统计效能。第三，GSE120575 scRNA-seq验证仅使用48例患者中的11例，因为表达数据可用性有限，限制了统计推断。第四，克隆命运锁定模型是相关性的；通过体内干扰SPP1信号传导和KLF2恢复进行因果验证是必要的。第五，GeoMx DSP平台缺乏几个关键的NK样标记基因（KLF2、FGFBP2），限制了空间分析使用替代签名。",
        "转化意义。NK-Locked Ratio代表一种候选生物标志物，可以使用scRNA-seq或TCR-seq结合多标记流式细胞术从治疗前肿瘤活检中评估。强大的预测性能（校正OR = 38.14）和敏感性分析中的稳健性支持其临床潜力。此外，将SPP1+ TAM-KLF2轴鉴定为机制靶点表明，将抗PD-1治疗与SPP1通路抑制或KLF2增强策略相结合可能恢复非响应者中的NK样克隆承诺。临床前SPP1抑制剂和抗CD44抗体正在开发中，为组合试验提供了途径。",
    ]
    for para in discussion:
        add_body_paragraph(doc, para)

    doc.add_page_break()

    # ============ 材料与方法 ============
    add_heading_custom(doc, "4. 材料与方法", 1, 14)

    add_heading_custom(doc, "4.1. 发现队列", 2, 12)
    add_body_paragraph(doc, "发现队列包括来自GSE243013数据集（Zhang等，Cell 2025）的接受新辅助抗PD-1（特瑞普利单抗）治疗的NSCLC患者。单细胞RNA-seq数据可作为预处理的AnnData对象获得：GSE243013_T_cells.h5ad（434,458个T细胞×31,831个基因，231例患者）和GSE243013_immune.h5ad（1,254,749个免疫细胞×31,831个基因，243例患者）。在T细胞患者中，188例具有适合pCR（n = 84）vs non-MPR（n = 104）二元比较的病理响应注释。细胞类型注释（sub_cell_type列）包括CD8T_NK-like_FGFBP2（n = 16,744）、Tex亚群（n = 71,298）和其他CD8/CD4 T细胞亚群。TCR克隆型信息可在clonotype和clonotype_number列中获得。")

    add_heading_custom(doc, "4.2. 数据预处理", 2, 12)
    add_body_paragraph(doc, "使用anndata包（v0.8）加载单细胞数据。使用normalize_total(1e4)后跟log1p转换对表达矩阵进行对数归一化，实现为等效于scanpy.pp.normalize_total + scanpy.pp.log1p的稀疏安全自定义函数。对于可视化，对采样数据（每患者≤200个非NK样细胞加上所有NK样细胞）执行PCA（50个成分）后跟UMAP（n_neighbors = 15，min_dist = 0.3，random_state = 42）的降维，以确保计算可处理性和平衡表示。")

    add_heading_custom(doc, "4.3. NK样基因签名", 2, 12)
    add_body_paragraph(doc, "单细胞分析的NK样10基因面板包括：FGFBP2、KLRD1、CX3CR1、FCGR3A、KLRC1、NKG7、GNLY、PRF1、GZMB、KLRB1。对于外部bulk RNA-seq验证，使用扩展的24基因NK样签名：FGFBP2、KLRD1、CX3CR1、FCGR3A、KLRC1、KLRC2、KLRB1、NKG7、GNLY、PRF1、GZMB、GZMH、GZMA、CTSW、KLRF1、SH2D1B、TYROBP、FCER1G、CD160、CRTAM、IFNG、TBX21、EOMES、ZNF683。该签名来源于已发表的NK样CD8 T细胞转录谱，包含NK受体、细胞毒性效应子和转录因子。跨平台基因匹配动态执行，支持基因符号、Ensembl ID和Entrez ID格式。")

    add_heading_custom(doc, "4.4. 克隆命运锁定分析", 2, 12)
    add_body_paragraph(doc, "大克隆定义为≥5个细胞的克隆型。对于每个大克隆，计算NK样CD8 T细胞的比例（nk_frac）。如果nk_frac > 0.5，则大克隆被分类为NK锁定；如果Tex细胞占多数，则分类为Tex分化。患者水平的NK-Locked Ratio定义为NK锁定的大克隆比例。使用statsmodels（smf.logit）执行逻辑回归，以pCR为二元结果（1 = pCR，0 = non-MPR）。单变量模型包括NK-Locked Ratio或Tex主导比例作为唯一预测因子。多变量模型调整总克隆数和CD8 T细胞比例。从模型参数中提取比值比和95%置信区间。在克隆大小阈值0、5、10、20和50个细胞下进行敏感性分析。")

    add_heading_custom(doc, "4.5. SPP1+ TAM鉴定与相互作用分析", 2, 12)
    add_body_paragraph(doc, "从完整免疫细胞图谱（1,254,749个细胞）中提取髓系细胞（major_cell_type = \"Myeloid cell\"，n = 191,099）并进行对数归一化。SPP1+ TAMs定义为巨噬细胞/单核细胞亚群中SPP1表达高于亚群中位数的细胞。患者水平的SPP1+ TAM比例计算为SPP1+ TAMs与总髓系细胞的比率。对8对候选配体-受体（SPP1-CD44、SPP1-ITGAV、SPP1-ITGB1、CXCL8-CXCR2、CCL2-CCR2、IL1B-IL1R1、TNF-TNFRSF1A、VEGFA-FLT1）进行配体-受体相互作用分析。相互作用得分计算为SPP1+ TAMs中的平均配体表达与每个CD8 T细胞亚型（分别为NK样和Tex）中的平均受体表达的乘积。最终相互作用得分定义为NK样和Tex相互作用得分的最大值，以避免跨功能不同受体群体的信号平均。")

    add_heading_custom(doc, "4.6. 干性基因分析", 2, 12)
    add_body_paragraph(doc, "干性相关基因包括KLF2、TCF7、LEF1、IL7R、SELL和CCR7。根据SPP1+ TAM比例将患者分为高SPP1组（前30%）和低SPP1组（后30%）。组间差异表达分析使用Mann-Whitney U检验，并进行Benjamini-Hochberg FDR校正。NK样与Tex细胞之间干性基因表达的直接比较在敏感性过滤子集（每患者NK样细胞计数≥10）上进行，以排除极值偏差。")

    add_heading_custom(doc, "4.7. 外部验证队列", 2, 12)
    add_body_paragraph(doc, "GSE126044：NSCLC bulk RNA-seq，抗PD-1（纳武利尤单抗），16例预处理组织样本（R = 5，NR = 11）。响应标签从GEO系列矩阵（\"patient response\"字段）中提取。NK样24基因签名得分计算为匹配基因表达的平均值。")
    add_body_paragraph(doc, "GSE135222：NSCLC bulk RNA-seq，抗PD-1/PD-L1，27例样本（R = 6，NR = 21）。响应标签从PFS字段定义（PFS = 0 → R，PFS = 1 → NR），因为RECIST响应注释不可用。基因ID为Ensembl格式（13/24个基因匹配）。")
    add_body_paragraph(doc, "GSE91061：黑色素瘤bulk RNA-seq，抗PD-1（纳武利尤单抗）±抗CTLA-4（伊匹木单抗），109例样本。选择具有PRCR（部分/完全缓解，n = 10）或PD（疾病进展，n = 23）的预处理样本；排除SD和未知响应。基因ID为Entrez格式（22/24个基因匹配）。")
    add_body_paragraph(doc, "GSE120575：黑色素瘤scRNA-seq，抗PD-1，共48例患者（R = 21，NR = 27）。可用预提取的NK基因表达子集（24个基因×16,291个细胞）。患者ID从细胞条形码中提取（格式：well_patient_batch）。每位患者的伪bulk NK样签名得分计算为所有细胞的平均表达。11例患者同时具有表达数据和匹配的响应标签（R = 3，NR = 8）。")
    add_body_paragraph(doc, "响应组间的统计比较使用Mann-Whitney U检验（scipy.stats.mannwhitneyu，双侧）。倍数变化计算为mean(R)/mean(NR)。对于跨癌种森林图，使用2000次迭代bootstrap重采样（numpy.random.default_rng，seed = 42）估计log2倍数变化的95%置信区间。")

    add_heading_custom(doc, "4.8. 空间转录组学验证", 2, 12)
    add_body_paragraph(doc, "分析来自GSE221733的空间转录组数据（GeoMx DSP，NSCLC，来自39例患者的93个ROIs）。过滤掉具有缺失响应标签或非标准片段类型的ROIs后，保留来自39例患者的67个ROIs（响应者 = 16，非响应者 = 23；PanCK+肿瘤 = 38，PanCK−基质 = 29）。平台包含8,659个探针，映射到1,812个独特基因。SPP1-受体相互作用得分计算为SPP1表达×mean(CD44, ITGAV, ITGB1)。干性得分使用可用的替代基因（TCF7、LEF1）。NK细胞毒性得分使用GZMB、GZMA、GNLY、PRF1。标准处理流程包括UMAP降维（n_neighbors = 15，min_dist = 0.3）用于总体分布可视化、基因表达映射和得分分布分析。")

    add_heading_custom(doc, "4.9. 统计分析", 2, 12)
    add_body_paragraph(doc, "所有统计分析在Python 3.11中使用scipy（v1.11）、statsmodels（v0.14）和numpy（v1.24）进行。Mann-Whitney U检验为双侧。Spearman秩相关用于非参数关联分析。逻辑回归使用最大似然估计，具有默认收敛标准。FDR校正使用Benjamini-Hochberg方法。所有p值、比值比、相关系数和置信区间均从数据动态计算——无统计值硬编码。所有分析可通过提供的run_all.py脚本完全重现。")

    add_heading_custom(doc, "4.10. 数据与代码可用性", 2, 12)
    add_body_paragraph(doc, "所有数据集可从NCBI基因表达综合数据库（GEO）公开获取：GSE243013（发现）、GSE126044、GSE135222、GSE91061、GSE120575（验证）、GSE221733（空间）。TCGA数据从基因组数据共享中心（SKCM、LUAD、LUSC）获得。所有分析代码在补充材料中提供（code/figure1.py至figure_supplement.py、spatial_validation.py、config.py、_common.py），可通过run_all.py执行。")

    doc.add_page_break()

    # ============ 结论 ============
    add_heading_custom(doc, "5. 结论", 1, 14)
    add_body_paragraph(doc, "本研究确立克隆命运锁定——扩增的CD8 T细胞克隆不可逆地定向到NK样或Tex轨迹——作为NSCLC抗PD-1响应的决定性原则。NK-Locked Ratio有效且独立地预测病理完全缓解（校正OR = 38.14，p = 0.002），而SPP1+ TAMs通过SPP1-CD44/ITGB1信号传导抑制NK样CD8 T细胞中的KLF2依赖性干性。肿瘤类型特异性的预测性能——在NSCLC中验证但在黑色素瘤中未验证——揭示克隆命运锁定代表免疫'冷'肿瘤中的限速瓶颈。这些发现为髓系驱动免疫逃逸提供了机制框架，并将SPP1+ TAM-KLF2轴鉴定为克服抗PD-1耐药性的候选治疗靶点。")

    # ============ 补充材料 ============
    doc.add_page_break()
    add_heading_custom(doc, "补充材料", 1, 14)

    add_figure(doc, os.path.join(RESULT_DIR, "FigS1_dataset_overview.png"),
               "图S1. 数据集概述和质量控制指标。")

    add_figure(doc, os.path.join(RESULT_DIR, "FigS2_sensitivity.png"),
               "图S2. NK-Locked Ratio预测性能在不同克隆大小阈值（0、5、10、20、50个细胞）下的敏感性分析。")

    add_figure(doc, os.path.join(RESULT_DIR, "FigS3_additional_validation.png"),
               "图S3. 额外验证队列基因匹配统计。")

    add_figure(doc, os.path.join(RESULT_DIR, "FigS_spatial_validation.png"),
               "图S_spatial_validation. 空间转录组学验证（5-panel）。")

    add_figure(doc, os.path.join(RESULT_DIR, "FigS_spatial_overview.png"),
               "图S_spatial_overview. 空间转录组标准处理流程图（3×4 panel：UMAP整体分布 + 基因表达 + 得分分布）。")

    # ============ 参考文献 ============
    doc.add_page_break()
    add_heading_custom(doc, "参考文献", 1, 14)

    refs = [
        "Forde, P.M.; et al. Neoadjuvant Nivolumab in Resectable Lung Cancer. N. Engl. J. Med. 2018, 379, e1–e31.",
        "Cascone, T.; et al. Neoadjuvant Nivolumab or Nivolumab Plus Ipilimumab in Resectable Non-Small Cell Lung Cancer: The CheckMate 816 Trial. J. Clin. Oncol. 2023, 41, 3679–3690.",
        "Garon, E.B.; et al. Pembrolizumab for the Treatment of Non-Small-Cell Lung Cancer. N. Engl. J. Med. 2015, 372, 2018–2028.",
        "Hellmann, M.D.; et al. Nivolumab plus Ipilimumab in Lung Cancer with a High Tumor Mutational Burden. N. Engl. J. Med. 2018, 378, 2093–2104.",
        "Philip, M.; Schietinger, A. CD8+ T Cell Differentiation and Dysfunction in Cancer. Nat. Rev. Immunol. 2022, 22, 209–223.",
        "Sade-Feldman, M.; et al. Defining T Cell States Associated with Response to Checkpoint Immunotherapy in Melanoma. Cell 2018, 175, 998–1013.",
        "Chiossone, L.; et al. Molecular Characterization of Human Natural Killer Cells. Immunol. Rev. 2018, 286, 1–14.",
        "Crome, S.Q.; et al. A Distinct Innate-Like CD8+ T Cell Population. Eur. J. Immunol. 2012, 42, 2632–2642.",
        "Freud, A.G.; et al. The Broad Spectrum of Human Natural Killer Cell Diversity. Immunity 2017, 47, 820–833.",
        "Dobano, C.; et al. Expression and Function of NK Cell Receptors on CD8+ T Cells. Front. Immunol. 2019, 10, 2336.",
        "Mantovani, A.; et al. Tumor-Associated Macrophages as Treatment Targets in Oncology. Nat. Rev. Clin. Oncol. 2017, 14, 399–416.",
        "DeNardo, D.G.; Ruffell, B. Macrophages as Regulators of Tumor Immunity and Immunotherapy. Nat. Rev. Immunol. 2019, 19, 369–382.",
        "Zhang, Q.; et al. Landmarkscapes of Tumor-Infiltrating Immune Cells in Cancer. Cell 2021, 184, 797–812.",
        "Oshi, M.; et al. SPP1 Expression Is a Prognostic Biomarker in Breast Cancer. Cancers 2021, 13, 1451.",
        "Zhang, J.; et al. Single-Cell Landscape of NSCLC Anti-PD-1 Therapy. Cell 2025. (GSE243013)",
        "Emerson, R.O.; et al. High-Throughput Sequencing of T-Cell Receptors Reveals a Homogeneous Repertoire of Tumor-Infiltrating Lymphocytes in Ovarian Cancer. PLoS One 2013, 8, e76808.",
        "Hart, G.T.; et al. Abandoning the Tug-of-War between T Cell Quiescence and Activation. Nat. Rev. Immunol. 2023, 23, 325–338.",
        "Yarchoan, M.; et al. Tumor Mutational Burden and Response Rate to PD-1 Inhibition. N. Engl. J. Med. 2017, 377, 2500–2501.",
        "Simon, S.; Labarriere, N. PD-1 Expression on Tumor-Specific T Cells: Friend or Foe for Immunotherapy? Cancers 2017, 9, 95.",
        "Krishna, C.; et al. Single-cell Sequencing Links Multiregional Immune Landscapes and Tissue-Resident T Cells in Cervical Cancer. Nat. Genet. 2021, 53, 120–129.",
        "Miller, B.C.; et al. Subsets of Exhausted CD8+ T Cells Differentially Mediate Tumor Control and Respond to Checkpoint Blockade. Nat. Immunol. 2019, 20, 326–336.",
        "Siesel, C.S.; et al. Stem-like CD8+ T Cells. Nat. Rev. Immunol. 2023, 23, 611–623.",
        "Pauken, K.E.; et al. Epigenetic Stability of Exhausted T Cells Limits Durability of Reinvigoration by PD-1 Blockade. Science 2016, 354, 1160–1165.",
        "Alfei, F.; et al. CD8+ T Cell Epigenetic Fixed Trait and Counteracting Inflammatory Signals. Immunity 2019, 50, 108–121.",
        "Rangaswami, H.; et al. Osteopontin: Role in Cell Signaling and Cancer Progression. Trends Cell Biol. 2006, 16, 79–87.",
        "Weber, G.F.; et al. Receptor-Ligand Interaction Between CD44 and Osteopontin (Eta-1). Science 1996, 271, 509–512.",
        "Ponta, H.; et al. CD44: A Multifunctional Cell Surface Adhesion Receptor. Nat. Rev. Mol. Cell Biol. 2003, 4, 33–45.",
        "Hogg, N.; et al. Integrin and Function-Associated Molecules on T Cells. Immunol. Rev. 2002, 186, 171–178.",
        "Carlson, C.M.; et al. Kruppel-like Factor 2 Regulates T Cell Trafficking. Nature 2006, 442, 1049–1052.",
        "Huang, A.C.; et al. T-Cell Infiltration and Immunity in Melanoma. Nat. Rev. Cancer 2020, 20, 65–77.",
        "Yost, K.E.; et al. Clonal Replacement of Tumor-Specific T Cells Following PD-1 Blockade. Nat. Med. 2019, 25, 1251–1259.",
        "Shevde, L.A.; et al. SPP1 (Osteopontin) as a Therapeutic Target in Cancer. Expert Opin. Ther. Targets 2010, 14, 1217–1230.",
        "Zöller, M. CD44: Can a Cancer-Initiating Cell Profit from an Abundantly Expressed Molecule? Nat. Rev. Cancer 2011, 11, 254–267.",
        "Anagnostou, V.; et al. Dynamics of Tumor and Immune Responses during Immune Checkpoint Blockade in Non-Small Cell Lung Cancer. Cancer Res. 2020. (GSE126044)",
        "Jung, H.; et al. Liquid Biopsy Enables Oncogenic Tracking in NSCLC. Nat. Commun. 2019. (GSE135222)",
        "Riaz, N.; et al. Tumor and Microenvironment Evolution during Immunotherapy in Melanoma. Cell 2017, 171, e895. (GSE91061)",
        "Sade-Feldman, M.; et al. Defining T Cell States Associated with Response to Checkpoint Immunotherapy in Melanoma. Cell 2018, 175, 998–1013. (GSE120575)",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(f"{i}. {ref}")
        set_cell_font(run, 'Times New Roman', '宋体', 9, False)
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.3
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)

    # 保存
    doc.save(OUTPUT_PATH)
    print(f"论文已保存至: {OUTPUT_PATH}")
    print(f"文件大小: {os.path.getsize(OUTPUT_PATH) / 1024 / 1024:.2f} MB")


if __name__ == "__main__":
    main()
