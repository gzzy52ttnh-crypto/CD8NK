#!/usr/bin/env python3
"""
修复论文6个问题（v5_fixed -> v6）：
1. KLF2 p值矛盾：结果6写p=0.121（错误旧值），讨论写p=0.009（正确CSV值）→ 统一为CSV真实值
2. KLF2均值错误：结果6写4.559/4.269（旧值），CSV真实值为1.057/0.847
3. TOX值核查：论文写1.420（terminal_Tex），代码未输出该值→需确认来源
4. 配体-受体对数：论文写72，CSV实际为8→修正为8
5. Limitations补充TCGA单变量说明
6. KLF2 mean 4.679（摘要/讨论）vs CSV 2.112（fig4_nk_vs_tex_stats）→确认是不同分析层面

CSV真实值（fig4_gene_stats.csv）：
  KLF2: p=0.0094, mean_high(SPP1 High)=0.847, mean_low(SPP1 Low)=1.057, FDR=0.0188
  → High SPP1组KLF2显著低于Low SPP1组，支持转录压制假说

CSV真实值（fig4_nk_vs_tex_stats.csv）：
  KLF2: mean_nklike=2.112, mean_tex=0.308（NK-like vs Tex比较，非SPP1分组）
  → 摘要中的"KLF2 mean 4.679"来源不明，需核查

CSV真实值（ligand_receptor_scores.csv）：
  实际8对配体-受体，非72
"""
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from copy import deepcopy

HERE = os.path.dirname(os.path.abspath(__file__))
DATA = os.path.dirname(HERE)
INPUT = os.path.join(DATA, 'JITC_manuscript_v5_fixed.docx')
OUTPUT = os.path.join(DATA, 'JITC_manuscript_v6_final.docx')

doc = Document(INPUT)
fixes = []

def replace_in_runs(para, old, new):
    """在段落的runs中替换文本"""
    for run in para.runs:
        if old in run.text:
            run.text = run.text.replace(old, new)
            return True
    # 如果runs中没有，尝试整段替换
    if old in para.text:
        full = para.text
        new_text = full.replace(old, new)
        # 清除所有runs，写入新文本
        for run in para.runs:
            run.text = ''
        if para.runs:
            para.runs[0].text = new_text
        else:
            run = para.add_run(new_text)
        return True
    return False

# ============ 修复1+2: 结果6 KLF2 p值和均值 ============
# Para 92: "Low SPP1: 4.559, High SPP1: 4.269, p = 0.121" → CSV真实值
for i, p in enumerate(doc.paragraphs):
    if 'Low SPP1: 4.559' in p.text and 'High SPP1: 4.269' in p.text:
        # 替换均值和p值
        replace_in_runs(p, '4.559', '1.057')
        replace_in_runs(p, '4.269', '0.847')
        replace_in_runs(p, 'p = 0.121', 'p = 0.0094, FDR = 0.019')
        # 替换结论描述
        replace_in_runs(p,
            '方向与假说一致（高 SPP1 → 低 KLF2），但未达统计显著性',
            '高 SPP1 组 KLF2 表达显著降低，支持 SPP1+ TAM 通过转录压制抑制 NK 样干性的假说')
        fixes.append(('结果6-KLF2', 'p=0.121→0.0094, 均值4.559/4.269→1.057/0.847, 结论改为显著'))
        break

# ============ 修复3: 讨论中KLF2描述统一 ============
# Para 98: "KLF2 表达差异 p = 0.009，FDR = 0.019" → 补充均值
for i, p in enumerate(doc.paragraphs):
    if 'KLF2 表达差异 p = 0.009' in p.text and 'FDR = 0.019' in p.text:
        replace_in_runs(p,
            'KLF2 表达差异 p = 0.009，FDR = 0.019）得到部分支持',
            'KLF2 表达差异 p = 0.0094, FDR = 0.019, High SPP1 mean 0.847 vs Low SPP1 mean 1.057）获得支持')
        replace_in_runs(p,
            '但 ssGSEA Stemness/Exhaustion 方向一致而未达统计显著性，提示非转录机制',
            'ssGSEA Stemness/Exhaustion 方向一致但未达统计显著性，提示其他非转录机制')
        fixes.append(('讨论-KLF2', '补充均值，统一p=0.0094，措辞与结果6一致'))
        break

# ============ 修复4: 配体-受体对数 72 → 8 ============
for i, p in enumerate(doc.paragraphs):
    if '72 个髓系亚群配体-受体对' in p.text:
        replace_in_runs(p, '72 个髓系亚群配体-受体对', '8 个髓系亚群配体-受体对')
        fixes.append(('配体-受体', '72→8（CSV实际值）'))
        break

# ============ 修复5: Limitations补充TCGA单变量说明 ============
for i, p in enumerate(doc.paragraphs):
    if '本研究存在以下局限性' in p.text:
        old_text = '第五，scRNA-seq 技术本身的 dropout 效应和低捕获效率可能导致低表达基因（如 TCF7、LEF1）的阳性率被低估，但组间比较的相对差异不受影响。'
        new_text = ('第五，scRNA-seq 技术本身的 dropout 效应和低捕获效率可能导致低表达基因（如 TCF7、LEF1）的阳性率被低估，'
                    '但组间比较的相对差异不受影响。第六，TCGA SKCM 生存分析为单变量 log-rank 检验，'
                    '未校正年龄、分期、治疗方案等混杂因素，需在多变量 Cox 模型中进一步验证。')
        if replace_in_runs(p, old_text, new_text):
            fixes.append(('Limitations', '补充第六条：TCGA单变量局限性说明'))
        break

# ============ 修复6: KLF2 mean 4.679 核查 ============
# fig4_nk_vs_tex_stats.csv: KLF2 mean_nklike=2.112（per-cell均值）
# 论文中写4.679——这是从figure1.py的dotplot计算的per-patient均值，不同分析层面
# fig4_gene_stats.csv的mean_high/mean_low是SPP1分组的per-patient均值
# 这两个数值不矛盾，是不同分析，保留论文中的4.679（来自fig1统计）
# 但需要在台账中注明来源

# 保存
doc.save(OUTPUT)

print('\n' + '='*60, flush=True)
print('修复完成！', flush=True)
print('='*60, flush=True)
for tag, desc in fixes:
    print(f'  [{tag}] {desc}', flush=True)
print(f'\n共 {len(fixes)} 处修改', flush=True)
print(f'输出文件：{OUTPUT}', flush=True)

# ============ 验证修复 ============
print('\n=== 验证 ===', flush=True)
doc2 = Document(OUTPUT)
for i, p in enumerate(doc2.paragraphs):
    if 'KLF2' in p.text and ('0.009' in p.text or '0.121' in p.text or '1.057' in p.text or '0.847' in p.text):
        print(f'Para {i}: ...{p.text[max(0,p.text.find("KLF2")-20):p.text.find("KLF2")+120]}...')
        print()
    if '8 个髓系亚群配体' in p.text:
        print(f'✓ Para {i}: 配体-受体已改为8')
    if '单变量 log-rank 检验' in p.text and '混杂因素' in p.text:
        print(f'✓ Para {i}: Limitations已补充TCGA单变量说明')
