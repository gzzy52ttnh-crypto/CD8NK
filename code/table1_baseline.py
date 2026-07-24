#!/usr/bin/env python3
"""
Step 4 — Table 1: 患者基线特征表 (Baseline Characteristics)
  从 GSE243013_immune.h5ad obs 提取患者级元数据
  按 pCR / MPR / non-MPR 三列分层
  组间比较：连续变量 ANOVA / Kruskal-Wallis，分类变量 χ² / Fisher
  输出：CSV + Word 表格
"""
import os, sys
import numpy as np
import pandas as pd
import scanpy as sc
import scipy.stats as ss
from scipy.stats import fisher_exact, chi2_contingency, kruskal, f_oneway
import warnings
warnings.filterwarnings('ignore')

HERE = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, HERE)
DATA = os.path.dirname(HERE)
RESULT = os.path.join(DATA, 'result')
os.makedirs(RESULT, exist_ok=True)

print('=== Loading GSE243013_immune.h5ad ===', flush=True)
adata = sc.read_h5ad(os.path.join(DATA, 'adata', 'GSE243013_immune.h5ad'))
# 去重到患者级别
obs = adata.obs.drop_duplicates('sampleID').copy()
obs['pathological_response'] = obs['pathological_response'].astype(str)
obs = obs[obs['pathological_response'].isin(['pCR', 'MPR', 'non-MPR'])]
print(f'n_patients = {len(obs)}  pCR={(obs["pathological_response"]=="pCR").sum()}  MPR={(obs["pathological_response"]=="MPR").sum()}  non-MPR={(obs["pathological_response"]=="non-MPR").sum()}', flush=True)

# ---- 数据清洗 ----
# Age: numeric
obs['age'] = pd.to_numeric(obs['age'], errors='coerce')

# Gender: M/F
obs['gender'] = obs['gender'].astype(str).str.strip()

# Smoking: Y/N -> Ever/Unknown, 且拼写错误 unknowm -> Unknown
obs['smoking'] = obs['smoking_history'].astype(str).str.strip()
obs['smoking'] = obs['smoking'].replace({'unknowm': 'Unknown', 'N': 'No', 'Y': 'Yes'})

# Cancer type
obs['cancer_type'] = obs['cancer_type'].astype(str).str.strip()

# Staging: 归并为 I/II/III/IV/Unknown
stage_map = {
    'IA2': 'I', 'IA3': 'I', 'IB': 'I',
    'IIA': 'II', 'IIB': 'II',
    'III': 'III', 'IIIA': 'III', 'IIIB': 'III', 'IIIC': 'III',
    'IVA': 'IV', 'IVB': 'IV',
    'unknown': 'Unknown',
}
obs['stage_group'] = obs['pre_treatment_staging'].astype(str).map(stage_map).fillna('Unknown')

# ---- 辅助函数：格式化数值 ----
def fmt_mean_sd(vals):
    vals = vals.dropna()
    return f'{vals.mean():.1f} ± {vals.std():.1f}'

def fmt_n_pct(n, total):
    return f'{int(n)} ({n/total*100:.1f}%)'

def fmt_range(vals):
    vals = vals.dropna()
    return f'{vals.min():.0f}–{vals.max():.0f}'

# ---- 构建 Table 1 ----
rows = []
resp_groups = ['pCR', 'MPR', 'non-MPR']
subs = {g: obs[obs['pathological_response'] == g] for g in resp_groups}
n_all = {g: len(subs[g]) for g in resp_groups}

# Header
rows.append(dict(
    Variable='n',
    pCR=str(n_all['pCR']),
    MPR=str(n_all['MPR']),
    non_MPR=str(n_all['non-MPR']),
    p_value='',
))

# ---- Age ----
age_vals = {g: subs[g]['age'] for g in resp_groups}
age_all = [age_vals[g].dropna() for g in resp_groups]
# Kruskal-Wallis (非正态友好)
kw_stat, kw_p = kruskal(*age_all)
rows.append(dict(
    Variable='Age, years, mean ± SD',
    pCR=fmt_mean_sd(age_vals['pCR']),
    MPR=fmt_mean_sd(age_vals['MPR']),
    non_MPR=fmt_mean_sd(age_vals['non-MPR']),
    p_value=f'{kw_p:.3f}',
))
rows.append(dict(
    Variable='  Range',
    pCR=fmt_range(age_vals['pCR']),
    MPR=fmt_range(age_vals['MPR']),
    non_MPR=fmt_range(age_vals['non-MPR']),
    p_value='',
))

# ---- Gender ----
rows.append(dict(Variable='Gender, n (%)', pCR='', MPR='', non_MPR='', p_value=''))
for cat in ['M', 'F']:
    counts = {g: (subs[g]['gender'] == cat).sum() for g in resp_groups}
    # Fisher 精确检验（2×3 表）——用 chi2 若期望频数足够
    ct = pd.crosstab(obs['gender'], obs['pathological_response'])
    try:
        _, p_gender, _, _ = chi2_contingency(ct[resp_groups])
        p_str = f'{p_gender:.3f}'
    except ValueError:
        _, p_gender = fisher_exact(ct[['pCR', 'non-MPR']])
        p_str = f'{p_gender:.3f}*'
    rows.append(dict(
        Variable=f'  {cat}',
        pCR=fmt_n_pct(counts['pCR'], n_all['pCR']),
        MPR=fmt_n_pct(counts['MPR'], n_all['MPR']),
        non_MPR=fmt_n_pct(counts['non-MPR'], n_all['non-MPR']),
        p_value=p_str if cat == 'M' else '',
    ))

# ---- Smoking ----
rows.append(dict(Variable='Smoking history, n (%)', pCR='', MPR='', non_MPR='', p_value=''))
for cat in ['Yes', 'No', 'Unknown']:
    counts = {g: (subs[g]['smoking'] == cat).sum() for g in resp_groups}
    ct = pd.crosstab(obs['smoking'], obs['pathological_response'])
    try:
        _, p_smoke, _, _ = chi2_contingency(ct[resp_groups])
        p_str = f'{p_smoke:.3f}'
    except ValueError:
        _, p_smoke = fisher_exact(ct[['pCR', 'non-MPR']])
        p_str = f'{p_smoke:.3f}*'
    rows.append(dict(
        Variable=f'  {cat}',
        pCR=fmt_n_pct(counts['pCR'], n_all['pCR']),
        MPR=fmt_n_pct(counts['MPR'], n_all['MPR']),
        non_MPR=fmt_n_pct(counts['non-MPR'], n_all['non-MPR']),
        p_value=p_str if cat == 'Yes' else '',
    ))

# ---- Cancer type ----
rows.append(dict(Variable='Histology, n (%)', pCR='', MPR='', non_MPR='', p_value=''))
for cat in ['LUAD', 'LUSC']:
    counts = {g: (subs[g]['cancer_type'] == cat).sum() for g in resp_groups}
    ct = pd.crosstab(obs['cancer_type'], obs['pathological_response'])
    try:
        _, p_ct, _, _ = chi2_contingency(ct[resp_groups])
        p_str = f'{p_ct:.3f}'
    except ValueError:
        _, p_ct = fisher_exact(ct[['pCR', 'non-MPR']])
        p_str = f'{p_ct:.3f}*'
    rows.append(dict(
        Variable=f'  {cat}',
        pCR=fmt_n_pct(counts['pCR'], n_all['pCR']),
        MPR=fmt_n_pct(counts['MPR'], n_all['MPR']),
        non_MPR=fmt_n_pct(counts['non-MPR'], n_all['non-MPR']),
        p_value=p_str if cat == 'LUAD' else '',
    ))

# ---- Stage ----
rows.append(dict(Variable='Clinical stage, n (%)', pCR='', MPR='', non_MPR='', p_value=''))
stage_cats = ['I', 'II', 'III', 'IV', 'Unknown']
for cat in stage_cats:
    counts = {g: (subs[g]['stage_group'] == cat).sum() for g in resp_groups}
    ct = pd.crosstab(obs['stage_group'], obs['pathological_response'])
    try:
        _, p_stage, _, _ = chi2_contingency(ct[resp_groups])
        p_str = f'{p_stage:.3f}'
    except ValueError:
        _, p_stage = fisher_exact(ct[['pCR', 'non-MPR']])
        p_str = f'{p_stage:.3f}*'
    rows.append(dict(
        Variable=f'  Stage {cat}',
        pCR=fmt_n_pct(counts['pCR'], n_all['pCR']),
        MPR=fmt_n_pct(counts['MPR'], n_all['MPR']),
        non_MPR=fmt_n_pct(counts['non-MPR'], n_all['non-MPR']),
        p_value=p_str if cat == 'I' else '',
    ))

# ---- Anti-PD1 therapy ----
# 统计接受 vs 未接受 anti-PD1（必须在 subs 创建前添加到 obs）
obs['received_anti_pd1'] = (~obs['anti-PD1_therapy'].astype(str).isin(['No', 'unknowm'])).astype(int)
# 重新创建 subs（因为加了新列）
subs = {g: obs[obs['pathological_response'] == g] for g in resp_groups}
rows.append(dict(Variable='Received anti-PD1 therapy, n (%)', pCR='', MPR='', non_MPR='', p_value=''))
for cat in [1, 0]:
    label = 'Yes' if cat == 1 else 'No'
    counts = {g: (subs[g]['received_anti_pd1'] == cat).sum() for g in resp_groups}
    ct = pd.crosstab(obs['received_anti_pd1'], obs['pathological_response'])
    try:
        _, p_pd1, _, _ = chi2_contingency(ct[resp_groups])
        p_str = f'{p_pd1:.3f}'
    except ValueError:
        _, p_pd1 = fisher_exact(ct[['pCR', 'non-MPR']])
        p_str = f'{p_pd1:.3f}*'
    rows.append(dict(
        Variable=f'  {label}',
        pCR=fmt_n_pct(counts['pCR'], n_all['pCR']),
        MPR=fmt_n_pct(counts['MPR'], n_all['MPR']),
        non_MPR=fmt_n_pct(counts['non-MPR'], n_all['non-MPR']),
        p_value=p_str if cat == 1 else '',
    ))

df_table = pd.DataFrame(rows)
csv_path = os.path.join(RESULT, 'table1_baseline.csv')
df_table.to_csv(csv_path, index=False)
print(f'Saved CSV -> {csv_path}', flush=True)

# ---- 生成 Word 表格 ----
try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print('python-docx not installed; generating Word table via pip install python-docx...', flush=True)
    os.system('/opt/anaconda3/envs/scanpy2/bin/pip install python-docx -q')
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn

doc = Document()
# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('Table 1. Baseline Characteristics of the Study Cohort')
run.bold = True
run.font.size = Pt(12)
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Table
table = doc.add_table(rows=1, cols=5)
table.style = 'Light Grid Accent 1'
hdr = table.rows[0].cells
headers = ['Characteristic', 'pCR (n={})'.format(n_all['pCR']), 'MPR (n={})'.format(n_all['MPR']),
           'non-MPR (n={})'.format(n_all['non-MPR']), 'p-value']
for i, h in enumerate(headers):
    p = hdr[i].paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.size = Pt(10)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# Fill rows
for _, row in df_table.iterrows():
    cells = table.add_row().cells
    for i, col in enumerate(['Variable', 'pCR', 'MPR', 'non_MPR', 'p_value']):
        val = str(row[col])
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(val)
        r.font.size = Pt(10)
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        # Bold for main categories (not indented)
        if i == 0 and not val.startswith('  '):
            r.bold = True

# Footnote
fn = doc.add_paragraph()
fn.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = fn.add_run('Note: Data are presented as mean ± standard deviation (SD) for continuous variables and as n (%) for categorical variables. p-values were calculated using the Kruskal-Wallis test for age and the chi-square test (or Fisher exact test where appropriate) for categorical variables.')
r.italic = True
r.font.size = Pt(9)
r.font.name = 'Times New Roman'
r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

fn2 = doc.add_paragraph()
fn2.alignment = WD_ALIGN_PARAGRAPH.LEFT
r2 = fn2.add_run('Abbreviations: pCR = pathologic complete response; MPR = major pathologic response; non-MPR = non-major pathologic response; LUAD = lung adenocarcinoma; LUSC = lung squamous cell carcinoma.')
r2.italic = True
r2.font.size = Pt(9)
r2.font.name = 'Times New Roman'
r2._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

docx_path = os.path.join(RESULT, 'table1_baseline.docx')
doc.save(docx_path)
print(f'Saved Word -> {docx_path}', flush=True)

# ---- 打印摘要到 console ----
print('\n=== Table 1 Summary ===', flush=True)
print(df_table.to_string(index=False), flush=True)
print('\n=== Key disclosures ===', flush=True)
print(f'Overall pCR rate: {n_all["pCR"]}/{sum(n_all.values())} = {n_all["pCR"]/sum(n_all.values())*100:.1f}%', flush=True)
print(f'LUSC proportion: {(obs["cancer_type"]=="LUSC").sum()}/{len(obs)} = {(obs["cancer_type"]=="LUSC").sum()/len(obs)*100:.1f}%', flush=True)

print('=== DONE Table 1 ===', flush=True)
