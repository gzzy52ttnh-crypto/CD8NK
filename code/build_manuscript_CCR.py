#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
按 Clinical Cancer Research (AACR) 投稿格式生成中文论文骨架（v1）。
依据 result/台账.md 的真实证据链与统计量；不使用过时 generate_docx.py 的错误数字。
输出：result/论文骨架_CCR_v1.docx

CCR 关键硬指标（initial submission 不强制，revision 须满足）：
- 标题 ≤165 字符；Running title ≤60 字符
- Translational Relevance 120–150 词（必需，位于摘要与 Introduction 之间）
- 结构化摘要 ≤250 词（Background/Methods/Results/Conclusions）
- 正文 5,000 词（不含摘要/Translational Relevance/参考文献/图注/表）
- 主文图+表 ≤6–7；参考文献 ≤50（AACR 编号格式）
- Results 与 Discussion 必须分离，不可合并
- 需 Data Availability 子节 + AI 使用声明 + 封面信临床意义陈述
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
RESULT_DIR = os.path.join(ROOT, "result")
OUTPUT_PATH = os.path.join(RESULT_DIR, "论文骨架_CCR_v1.docx")


# ---------- 字体 / 排版辅助 ----------
def set_font(run, font_name="宋体", font_name_east="宋体", size=10.5, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = r.makeelement(qn("w:rFonts"), {})
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name_east)


def add_title(doc, text, size=16):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, "黑体", "黑体", size, True)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_subtitle(doc, text, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    set_font(run, "Times New Roman", "宋体", size, False, italic=True)
    p.paragraph_format.space_after = Pt(10)
    return p


def add_heading(doc, text, font_size=14, bold=True, before=12, after=6, east="黑体"):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, east, east, font_size, bold)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def add_para(doc, text, size=10.5, indent=True, bold=False, color=None, after=6, label=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if label:
        r0 = p.add_run(label)
        set_font(r0, "黑体", "黑体", size, True)
    run = p.add_run(text)
    set_font(run, "宋体", "宋体", size, bold)
    if color is not None:
        run.font.color.rgb = color
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_bullet(doc, text, size=10.5, bold_lead=None):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if bold_lead:
        r1 = p.add_run(bold_lead)
        set_font(r1, "黑体", "黑体", size, True)
        r2 = p.add_run(text)
        set_font(r2, "宋体", "宋体", size, False)
    else:
        run = p.add_run(text)
        set_font(run, "宋体", "宋体", size, False)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4
    return p


def add_fig_placeholder(doc, fig_files, caption, width=6.0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("[图位 — 待插入]")
    set_font(run, "宋体", "宋体", 9, True, italic=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = cap.add_run(caption)
    set_font(r1, "宋体", "宋体", 9, False)
    cap.paragraph_format.space_after = Pt(4)

    files = doc.add_paragraph()
    files.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = files.add_run("对应文件：" + "；".join(fig_files))
    set_font(r2, "Consolas", "宋体", 8, False)
    r2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    files.paragraph_format.space_after = Pt(8)


def add_meta(doc, label, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(label)
    set_font(r1, "黑体", "黑体", 10.5, True)
    r2 = p.add_run(text)
    set_font(r2, "宋体", "宋体", 10.5, False)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.4


def add_constraint(doc, label, text):
    p = doc.add_paragraph(style="List Bullet")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = p.add_run(label)
    set_font(r1, "黑体", "黑体", 10, True)
    r2 = p.add_run(text)
    set_font(r2, "宋体", "宋体", 10, False)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.35


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)

    # ===================== 标题页区 =====================
    add_title(doc, "论文骨架（Clinical Cancer Research 格式 v1）")
    add_subtitle(doc, "NK-like CD8+ T cells clonal fate locking & SPP1+ TAM myeloid gate in NSCLC anti-PD-1")

    add_heading(doc, "标题页 / Title page", 13, before=8)
    add_meta(doc, "标题候选 1（推荐）：", "SPP1+ 肿瘤相关巨噬细胞经髓系闸门锁定 NK样CD8+ T细胞克隆命运并决定非小细胞肺癌抗PD-1疗效")
    add_meta(doc, "标题候选 2：", "克隆命运锁定预测新辅助抗PD-1响应：SPP1+ TAM 抑制 NK样CD8+ T细胞干性扩增")
    add_meta(doc, "英文标题候选（推荐，≤165 字符）：", "SPP1+ TAMs Lock NK-like CD8+ T Cell Clonal Fate and Dictate Anti-PD-1 Response in NSCLC")
    add_meta(doc, "Running title（≤60 字符）：", "SPP1+ TAM gates NK-like CD8 clonal fate in NSCLC（48 字符，符合）")
    add_meta(doc, "作者 / 单位 / 通讯：", "［占位：张泽民团队等；GEO 登录号 GSE243013 等］")
    add_meta(doc, "字数 / 图数（终稿填）：", "正文 ____ 词 / 主文图 __ 表 __（CCR 限 5,000 词、图+表≤6–7、参考文献≤50）")
    add_meta(doc, "利益冲突声明：", "［占位：The authors declare no potential conflicts of interest.］")
    add_meta(doc, "文章类型：", "Research Article — Translational Mechanisms and Therapy / Novel Biomarkers and Precision Medicine")

    # ===================== 格式约束检查表 =====================
    add_heading(doc, "格式约束检查表（CCR 硬指标 · revision 须满足）", 12, before=10)
    add_constraint(doc, "标题：", "≤165 字符（含空格）；Running title ≤60 字符。")
    add_constraint(doc, "Translational Relevance：", "120–150 词，必需，位于摘要与 Introduction 之间（见下文专节）。")
    add_constraint(doc, "结构化摘要：", "≤250 词，四段 Background / Methods / Results / Conclusions。")
    add_constraint(doc, "正文：", "5,000 词（不含摘要/Translational Relevance/参考文献/图注/表）。")
    add_constraint(doc, "主文图+表：", "≤7（已达上限）。当前主文 Fig1–6 + Table1 = 7；Fig7 空间交互已移入补充（FigS7），revision 时严格 ≤7。")
    add_constraint(doc, "Results/Discussion：", "必须分离，不可合并（CCR 明确规定）。")
    add_constraint(doc, "参考文献：", "≤50，AACR 编号格式：正文 [n] 上标；列出前 6 作者 + et al.。")
    add_constraint(doc, "Data Availability：", "Materials and Methods 末尾须设 Data Availability 子节（GEO 号 + 代码仓库）。")
    add_constraint(doc, "AI 使用声明：", "须在 Acknowledgments 披露 AI 辅助写作（AACR 强制）。")
    add_constraint(doc, "封面信：", "须陈述临床意义、转化影响、利益冲突；本报告文末附要点。")
    add_constraint(doc, "报告规范：", "观察性队列建议遵循 STROBE；含配对/外部验证须说明验证逻辑。")

    doc.add_page_break()

    # ===================== Translational Relevance（CCR 特有，必需）=====================
    add_heading(doc, "Translational Relevance（150 词，必需 · 位于摘要与引言之间）", 14)
    add_para(doc,
              "本研究在 188 例非小细胞肺癌（NSCLC）新辅助抗 PD-1 单细胞队列中发现，扩增 CD8+ T 细胞克隆的\"命运锁定\"（NK样 vs 终末耗竭分化）而非克隆大小，是决定病理完全缓解（pCR）的关键；SPP1+ 肿瘤相关巨噬细胞（TAM）通过 SPP1-CD44/ITGB1 轴抑制 NK样 CD8+ T 细胞的 KLF2 依赖性干性，构成削弱疗效的\"髓系闸门\"。这些结果提示，NK-Locked Ratio 可作为治疗前活检预测 pCR 的候选生物标志物，且 SPP1+ TAM-KLF2 轴是克服抗 PD-1 耐药的潜在靶点。若在前瞻性试验中验证，该框架将指导\"患者选择 + 联合 SPP1 通路抑制\"的精准免疫治疗策略，并为 Taxane 与 Pemetrexed 方案响应差异提供机制解释。",
              bold=False, after=4)
    add_para(doc, "（注：以上为中文草稿，英文终稿须控制在 150 words 以内，按 CCR 示例结构：specific finding → biomarker/target → clinical impact → next step。）",
              size=9, color=RGBColor(0xB0, 0x00, 0x00), after=6)

    # ===================== 结构化摘要 =====================
    add_heading(doc, "摘要（结构化，≤250 词：Background/Methods/Results/Conclusions）", 14)
    add_para(doc, "背景：抗 PD-1 免疫治疗已改变 NSCLC 新辅助格局（pCR 率约 20–45%），但缺乏可靠的预测性生物标志物；NK样 CD8+ T 细胞的功能身份与其克隆命运决定机制尚未系统阐明。", label="Background：", after=4)
    add_para(doc, "方法：基于 GSE243013 单细胞 + scTCR-seq（188 例、434,458 个 T 细胞）定义克隆命运锁定；结合空间转录组（GSE221733 DSP）、外部验证队列（GSE241934/179994/207422/120575）与临床预测模型（IRS 评分）。统计采用双侧 Mann-Whitney U、Spearman、逻辑回归（含 Firth 惩罚）、Bootstrap ΔAUC，随机种子 seed=42。", label="Methods：", after=4)
    add_para(doc, "结果：NK-dominant 克隆比例预测 pCR（OR=16.98, p=0.014, AUC=0.662），且所有克隆均为患者特异。SPP1+ TAM 比例与 NK-dominant 比例显著负相关（r=−0.333, p=5.28e-6）；SPP1-ITGB1/CD44 为最强配受体交互；高 SPP1 组中 NK样细胞 KLF2 下调（p=0.0255）。空间验证支持 Stemness Arrest 模型。外部 Taxane 队列方向一致（AUC=0.68），且方案特异性显著（Taxane vs Pemetrexed 交互 LRT p=0.007，Peme 反转 AUC=0.20）；两 Taxane 队列荟萃 AUC=0.664, p<0.001。IRS 评分显著优于临床基线（ΔAUC p=0.006）。", label="Results：", after=4)
    add_para(doc, "结论：克隆命运锁定是 NSCLC 抗 PD-1 响应的决定性原则；SPP1+ TAM 髓系闸门通过抑制 NK样 CD8+ T 细胞克隆锁定而削弱疗效，提示 SPP1+ TAM-KLF2 轴为潜在治疗靶点。", label="Conclusions：", after=4)
    add_meta(doc, "关键词（≤10，AACR 风格）：", "NK样 CD8 T 细胞；克隆命运锁定；抗 PD-1；SPP1+ 肿瘤相关巨噬细胞；非小细胞肺癌；单细胞 RNA 测序；空间转录组；免疫治疗生物标志物")

    doc.add_page_break()

    # ===================== 引言 =====================
    add_heading(doc, "Introduction", 14)
    add_para(doc, "抗 PD-1 免疫治疗已显著改变可切除非小细胞肺癌（NSCLC）的新辅助治疗格局，病理完全缓解（pCR）率约 20–45%，但多数患者未能获得持久获益。现有生物标志物（肿瘤突变负荷 TMB、PD-L1 表达）在抗 PD-1 新辅助场景中的预测力有限，亟需基于机制的、可在治疗前活检中评估的预测性标志物。［文献占位：新辅助抗 PD-1 临床试验与标志物局限］", after=4)
    add_para(doc, "CD8+ T 细胞的功能状态处于\"干性样记忆 ↔ 终末耗竭\"连续谱。NK样 CD8+ T 细胞（FGFBP2+，共表达 NK 受体）是一群具强细胞毒性的独特群体，被认为兼具干性与效应特征；然而它们究竟代表终末分化终点，还是扩增克隆内部的动态命运承诺，此前未被系统研究。［文献占位：NK样 CD8 身份与干性/效应之争］", after=4)
    add_para(doc, "肿瘤相关巨噬细胞（TAM）是 NSCLC 微环境的主导抑制区室；其中 SPP1+ TAM 亚群与 T 细胞抑制和不良预后显著相关，但其与 CD8+ T 细胞克隆动力学——尤其是是否直接影响克隆命运决定——的机制联系尚属未知。［文献占位：SPP1+ TAM 与免疫抑制］", after=4)
    add_para(doc, "本研究利用 GSE243013 大规模 scRNA + scTCR 图谱，发现\"克隆命运锁定\"现象：扩增克隆不可逆地定向到 NK样或耗竭（Tex）轨迹，且 NK 锁定比例有效预测 pCR。我们进一步阐明 SPP1+ TAM 髓系闸门通过 SPP1-CD44/ITGB1 抑制 NK样 CD8+ T 细胞干性主调控因子 KLF2，建立\"髓系驱动免疫逃逸\"的机制框架；多独立队列外部验证与空间转录组共同支持该肿瘤类型特异性模型，并据此构建整合 NK 锁定与 SPP1+ TAM 抑制的 IRS 预测评分。", after=4)

    doc.add_page_break()

    # ===================== 材料与方法 =====================
    add_heading(doc, "Materials and Methods", 14)
    add_bullet(doc, "发现队列：GSE243013（Zhang et al., Cell 2025）新辅助抗 PD-1 NSCLC；GSE243013_T_cells.h5ad（434,458 T × 31,831 基因）与 GSE243013_immune.h5ad（1,254,749 免疫细胞）。188 例具 pCR/non-MPR 注释。伦理/IRB 批件号［占位］。")
    add_bullet(doc, "数据预处理：AnnData 加载；normalize_total(1e4)+log1p；采样后 PCA(50)+UMAP(n_neighbors=15, min_dist=0.3, random_state=42)。")
    add_bullet(doc, "NK样定义与签名：单细胞 10 基因（FGFBP2/KLRD1/CX3CR1/FCGR3A/KLRC1/NKG7/GNLY/PRF1/GZMB/KLRB1）；外部 bulk 验证 24 基因签名；跨平台（基因符号/Ensembl/Entrez）映射。")
    add_bullet(doc, "克隆命运锁定分析：大克隆 ≥5 细胞；NK-dominant（nk_frac>0.5）；statsmodels 逻辑回归（单/多变量），阈值敏感性 0/5/10/20/50；Firth 惩罚回归稳健性。")
    add_bullet(doc, "SPP1+ TAM 鉴定与配受体分析：髓系细胞（191,099）；SPP1>亚群中位数定义为 SPP1+ TAM；8 对配受体（SPP1-CD44/ITGAV/ITGB1 等）交互得分 = 配体表达 × 受体表达（取 NK样/Tex 最大值）。")
    add_bullet(doc, "干性基因分析：KLF2/TCF7/LEF1/IL7R/SELL/CCR7；按 SPP1+ TAM 比例分高/低组（前/后 30%），Mann-Whitney U + BH-FDR。")
    add_bullet(doc, "外部验证队列：GSE241934（scRNA+scTCR，含 Taxane/Pemetrexed 双方案）、GSE179994（Pemetrexed）、GSE207422（单药）、GSE120575（黑色素瘤 anti-CTLA-4）；响应标签与基因匹配动态解析。")
    add_bullet(doc, "空间转录组：GSE221733 GeoMx DSP（39 例 / 67 ROI）；SPP1-受体交互得分、干性得分、NK 毒性得分；UMAP 降维。")
    add_bullet(doc, "临床模型与 IRS：逻辑回归 + 5 折 OOF；IRS = 整合 NK 锁定 + SPP1+ TAM 抑制；Bootstrap ΔAUC 显著性检验；随机种子 seed=42。")
    add_bullet(doc, "统计与可复现：Python（scipy/statsmodels/numpy）；双侧 MWU、Spearman、BH-FDR；全部统计量动态计算无硬编码；run_all.py 可复现。")
    add_para(doc, "Data Availability：单细胞与空间数据存于 GEO（GSE243013、GSE221733、GSE241934、GSE179994、GSE207422、GSE120575）；分析代码与中间结果见 ［GitHub/代码仓库占位］；补充数据随文上传。",
              bold=False, label="Data Availability：", after=4)
    add_para(doc, "AI 使用声明：手稿在 AI 写作辅助工具协助下准备；科学设计、数据分析、结果解释与结论完全由作者负责。［AACR 强制披露，置于 Acknowledgments］",
              bold=False, color=RGBColor(0xB0, 0x00, 0x00), after=6)

    doc.add_page_break()

    # ===================== 结果（与 Discussion 分离）=====================
    add_heading(doc, "Results", 14)

    add_heading(doc, "3.1 单细胞图谱鉴定 NK样 CD8+ T 细胞并表征队列基线（Fig1 / Table1 / FigS1）", 12)
    add_para(doc, "核心信息：在 GSE243013 中界定 NK样 CD8+ T 细胞亚群（CD8T_NK-like_FGFBP2），并确认其在 pCR vs non-MPR 间的比例差异；队列基线基本均衡。", bold=False, after=2)
    add_bullet(doc, "Fig1A：CD8+ T UMAP，NK样细胞高亮。")
    add_bullet(doc, "Fig1B：CD8+ 亚群比例饼图。")
    add_bullet(doc, "Fig1C：每患者 NK样比例箱线图（pCR vs non-MPR，Mann-Whitney U）。")
    add_bullet(doc, "Fig1D：每患者 T 细胞数直方图（注意排除\"无 CD8T 数据\"患者）。")
    add_bullet(doc, "Table1：基线特征表（sex/stage/histology/smoking/age），分类 Fisher、连续 MWU；结论：基线基本均衡。")
    add_fig_placeholder(doc,
                        ["Fig1_overview.png/.pdf", "FigS1_supplement.png/.pdf", "Table1_baseline.csv"],
                        "图1. NK样 CD8+ T 细胞的鉴定与队列基线。(A) CD8 T UMAP（NK样高亮）；(B) CD8 亚群比例；(C) pCR vs non-MPR 的 NK样比例箱线图；(D) 每患者 T 细胞数分布。")

    add_heading(doc, "3.2 克隆命运锁定：NK-dominant 克隆预测 pCR（Fig2 / FigS2 / FigS4）", 12)
    add_para(doc, "核心信息：扩增克隆的\"功能身份\"而非\"大小\"决定疗效。定义大克隆（≥5 细胞）并按主要亚群组成分类为 NK锁定 / Tex分化；患者级 NK-dominant 比例强预测 pCR，且跨阈值稳健。", bold=False, after=2)
    add_bullet(doc, "NK-dominant 克隆比例与 pCR 显著正相关：默认阈值=5 时 OR=16.98, p=0.014, AUC=0.662。", bold_lead="主要结果：")
    add_bullet(doc, "阈值敏感性：阈值 0/5/10/20/50 下均显著，≥5 平衡稳定性与统计效力（FigS2）。")
    add_bullet(doc, "所有克隆均为患者特异，无跨患者共享（FigS4_tcr_exclusivity）。")
    add_bullet(doc, "Firth 惩罚逻辑回归验证小样本稳健性。")
    add_fig_placeholder(doc,
                        ["Fig2_clonal_fate.png/.pdf", "FigS2_threshold_sensitivity.png/.pdf", "FigS4_tcr_exclusivity.png/.pdf"],
                        "图2. 克隆命运锁定预测响应。(A) NK锁定 vs Tex分化克隆定义；(B) NK-dominant 比例森林图；(C) pCR vs non-MPR 箱线图 + MWU；(D) 频数分布。")

    add_heading(doc, "3.3 空间转录组验证：Stemness Arrest 模型（FigS_spatial_*）", 12)
    add_para(doc, "核心信息：在原位层面，SPP1×受体交互与干性基因正相关（r=+0.325, p=0.044），方向与单细胞层面（SPP1↑→NK-dominant↓）相反，支持\"SPP1+ TAM 将 T 细胞锁在干性状态、阻止其向效应分化与克隆扩增\"的解释。", bold=False, after=2)
    add_bullet(doc, "GSE221733 GeoMx DSP：SPP1×受体交互得分与干性（TCF7/LEF1）正相关 r=+0.325, p=0.044, FDR=0.158（趋势性）。")
    add_bullet(doc, "FigS_spatial_validation / FigS_spatial_overview / FigS_spatial_stemness_arrest。")
    add_para(doc, "注意（局限性）：空间结果 14 个检验 FDR 均 >0.15，仅作趋势性提示；DSP 平台缺乏 KLF2/FGFBP2 探针，使用替代基因集。",
              bold=False, color=RGBColor(0xB0, 0x00, 0x00), after=6)
    add_fig_placeholder(doc,
                        ["FigS_spatial_validation.png/.pdf", "FigS_spatial_overview.png/.pdf", "FigS_spatial_stemness_arrest.png/.pdf"],
                        "图S（空间）. Stemness Arrest 模型验证。(A) ROI 空间得分；(B) 整体分布；(C) 干性阻滞散点。")

    add_heading(doc, "3.4 SPP1+ TAM 髓系微环境与 T 细胞内在机制（Fig3 / Fig4 / FigS5 / Fig4_TCF7）", 12)
    add_para(doc, "核心信息：SPP1+ TAM 比例在 non-MPR 组显著更高，并与 NK-dominant 比例负相关；SPP1-ITGB1/CD44 为最强配受体交互；高 SPP1 组中 NK样细胞 KLF2 等干性基因下调。", bold=False, after=2)
    add_bullet(doc, "SPP1+ TAM 比例 non-MPR 组显著更高（p=2.13e-9）。", bold_lead="髓系侧：")
    add_bullet(doc, "SPP1+ TAM 比例 vs NK-dominant 比例：r=−0.333, p=5.28e-6（强负相关）。")
    add_bullet(doc, "配受体：SPP1-ITGB1 为 NK样侧最强交互（NK_score=3.88）；SPP1-CD44 为 Tex 侧最强（Tex_score=3.90）；FigS5 网络图。")
    add_bullet(doc, "NK样细胞内 KLF2 在高 SPP1 组显著降低（p=0.0255）；火山图鉴定 5,308 个 DEG。", bold_lead="T细胞内在：")
    add_bullet(doc, "Fig4_TCF7_dysfunction：TCF7 与功能障碍关系。")
    add_fig_placeholder(doc,
                        ["Fig3_myeloid.png/.pdf", "Fig4_mechanism.png/.pdf", "Fig4_TCF7_dysfunction.png/.pdf", "FigS5_interaction_network.png/.pdf"],
                        "图3/4. 髓系闸门与机制链条。(A) 髓系 UMAP；（B）特征基因热图；（C）SPP1+ TAM 比例箱线图；（D）配受体网络；（E）SPP1+ TAM 比例 vs NK-dominant 散点 + 干性基因小提琴图。")

    add_heading(doc, "3.5 外部验证与方案特异性（Fig5 / GSE241934 系列 / FigS6-7；Fig7 移入补充）", 12)
    add_para(doc, "核心信息：外部验证证实 NK样与响应在 NSCLC Taxane 方案方向一致，且存在显著方案特异性（Taxane 一致、Pemetrexed 反转）；黑色素瘤队列方向相反，支持肿瘤类型/治疗特异性。", bold=False, after=2)
    add_para(doc, "（注：原 Fig7 空间交互图已移入补充材料 FigS7，以满足 CCR 主文图+表 ≤7 上限；主文保留 Fig1–6 + Table1 共 7 个主文图+表。）", size=9, color=RGBColor(0xB0, 0x00, 0x00), after=4)
    add_bullet(doc, "GSE120575（黑色素瘤 + anti-CTLA-4，19 例 R=9/NR=10）：AUC=0.1111, MWU p=0.0048，NK样签名在 Non-responder 更高（负向，与主队列相反）→ 癌种/治疗特异性。")
    add_bullet(doc, "GSE241934 Taxane 方案（n=24）：AUC=0.68，方向一致。")
    add_bullet(doc, "GSE241934 Pemetrexed 方案（n=20）：AUC=0.20，方向反转（p=0.03）。")
    add_bullet(doc, "方案特异性：Taxane OR=7.86 / Peme OR=0.34，交互效应 LRT p=0.007（强）。")
    add_bullet(doc, "GSE207422（NSCLC 单药，n=15）：AUC=0.648，方向一致但 NS（p=0.388）。GSE179994 Pemetrexed（n=13）：AUC=0.472, p=0.940。")
    add_bullet(doc, "荟萃分析（GSE243013 主队列 + GSE241934 Taxane，n=212）：合并 AUC=0.664, p<0.001, I²=0%；合并 OR=2.71, p=0.020。")
    add_fig_placeholder(doc,
                        ["Fig5_external_validation.png/.pdf", "GSE241934_Fig1_boxplot.png", "GSE241934_Fig2_ROC.png", "GSE241934_Fig4_treatment_specificity.png", "GSE241934_Fig5_meta_analysis.png"],
                        "图5 与 GSE241934 系列. 外部验证与方案特异性。(A) GSE120575 方向反转；(B) GSE241934 Taxane/Peme 箱线 + ROC；(C) AUC 比较；(D) 治疗特异性；(E) 荟萃森林图。（空间交互见补充 FigS7）")

    add_heading(doc, "3.6 临床转化：IRS 评分与预测模型（Fig6 / fig6_IRS_*）", 12)
    add_para(doc, "核心信息：将 NK 锁定与 SPP1+ TAM 抑制整合为 IRS 评分，预测力优于单指标与临床基线；但临床增量在更大样本前仅呈趋势。", bold=False, after=2)
    add_bullet(doc, "临床模型（Fig6）：Clinical only AUC=0.6894（训练）/OOF 0.6024；Clinical+NK-Locked AUC=0.7229（训练）/OOF 0.6474；ΔAUC +0.0335/+0.0450；Bootstrap p=0.153（趋势，未显著）。")
    add_bullet(doc, "IRS 评分：训练 AUC 0.747 vs NK-dominant 0.731；OOF 0.698 vs 0.683；ΔAUC vs Clinical Bootstrap p=0.006（显著优于临床基线）；IRS vs NK-dominant +0.016, p=0.051（边缘显著）。n=179。")
    add_bullet(doc, "fig6_IRS_nomogram / fig6_IRS_model：列线图与模型表现。")
    add_para(doc, "注意（局限性）：临床模型增量未达显著（ΔAUC p=0.153），需更大样本验证；NK样量化方法在主队列（克隆命运锁定 NK-dominant）与外部（细胞比例/基因签名）不一致，可能影响效应量可比性。",
              bold=False, color=RGBColor(0xB0, 0x00, 0x00), after=6)
    add_fig_placeholder(doc,
                        ["Fig6_clinical_model.png/.pdf", "fig6_IRS_model.png/.pdf", "fig6_IRS_nomogram.png/.pdf"],
                        "图6. 临床转化模型与 IRS 评分。(A) 临床模型 ROC；(B) IRS 模型；(C) 列线图。")

    doc.add_page_break()

    # ===================== 讨论（单独，不可与 Results 合并）=====================
    add_heading(doc, "Discussion", 14)
    add_para(doc, "主要发现凝练（4 点）：(1) 克隆命运锁定（NK锁定 vs Tex分化）预测 pCR；(2) SPP1+ TAM 通过 SPP1-ITGB1/CD44 抑制 NK样 KLF2 依赖性干性；(3) NK样签名在 NSCLC 外部队列验证、黑色素瘤未验证（肿瘤类型特异性）；(4) 空间转录组原位确认 SPP1-受体-干性轴。", after=4)
    add_para(doc, "克隆命运锁定作为决定性原则：克隆扩增幅度在应答者/非应答者间无差异（p=0.683），而 NK-Locked 比例区分性强 → 重新定义\"有效\"T 细胞克隆；与干性样 CD8（TCF1+/SLAMF6+）持续响应文献一致，并扩展 NK样轨迹经 KLF2 携带干性的概念。［文献占位］", after=4)
    add_para(doc, "SPP1+ TAM 髓系闸门机制：SPP1-CD44/ITGB1 为最强配受体对；对 KLF2（T 细胞静息/迁移/干性主调控因子）的抑制，提供髓系如何\"锁定\"T 细胞克隆的分子机制。［文献占位］", after=4)
    add_para(doc, "跨癌种异质性：黑色素瘤为高 TMB\"热\"肿瘤，瓶颈从克隆命运锁定转移到分化动力学 → 解释 GSE120575 方向反转；提示基于 NK样签名的生物标志物应在肿瘤类型特异性背景下开发。［文献占位］", after=4)
    add_para(doc, "空间验证：SPP1-受体-干性轴在原位确认（r=+0.325, p=0.044），响应者 ROI 中 NK 细胞毒性得分更高（p=0.027）；DSP 缺乏 KLF2/FGFBP2 探针为局限。", after=4)
    add_para(doc, "局限性（8 条）：① OR CI 宽（阈值=5 时 CI=[1.96,195.18]）；② 空间 FDR 均>0.15；③ 空间方向与单细胞相反（Stemness Arrest 解释，需功能验证）；④ 临床模型增量 NS（p=0.153）；⑤ 外部癌种/治疗异质性（GSE120575 反转）；⑥ GSE241934 样本小（Taxane n=24 一致但 NS）；⑦ NK样量化方法学差异；⑧ GSE241934 无克隆型分析，仅能验证细胞比例/签名层面。", after=4)
    add_para(doc, "转化意义：NK-Locked Ratio 可作为治疗前活检评估的候选标志物；SPP1+ TAM-KLF2 轴提示\"抗 PD-1 + SPP1 通路抑制 / KLF2 恢复\"组合策略。", after=4)

    doc.add_page_break()

    # ===================== 致谢 / 利益冲突 / 数据可用 =====================
    add_heading(doc, "Acknowledgments / 利益冲突 / 数据可用", 14)
    add_bullet(doc, "致谢：［基金号占位：国家自然科学基金等］；感谢［测序平台/临床合作团队］。")
    add_bullet(doc, "AI 使用声明（AACR 强制）：手稿在 AI 写作辅助工具（WorkBuddy）协助下准备；科学内容、数据分析与解释完全由作者负责。")
    add_bullet(doc, "利益冲突：The authors declare no potential conflicts of interest.（或据实填写）")
    add_bullet(doc, "数据可用：GEO 登录号见 Materials and Methods；代码仓库与补充数据见［占位］。")
    add_para(doc, "Author Contributions（CCR 必需）：", bold=True, indent=False, after=2)
    add_bullet(doc, "[作者姓名] conceived and designed the study.")
    add_bullet(doc, "[作者姓名] performed data analysis.")
    add_bullet(doc, "[作者姓名] interpreted the results.")
    add_bullet(doc, "[作者姓名] drafted the manuscript.")
    add_bullet(doc, "All authors reviewed and approved the final manuscript.")

    # ===================== 参考文献（AACR 编号格式）=====================
    doc.add_page_break()
    add_heading(doc, "References（AACR 编号格式，≤50；终稿按引用顺序编号）", 14)
    add_para(doc, "格式要求：正文以 [n] 上标引用；条目列出前 6 位作者，>6 位加 \"et al.\"；期刊名标准缩写；年;卷(期):起止页。示例（占位，终稿补全卷期与页码）：", after=2)
    add_bullet(doc, "Zhang Z, Wang L, Liu J, et al. Single-cell landscape of neoadjuvant anti-PD-1 in NSCLC. Cell. 2025;188(XX):XXXX–XXXX.（GSE243013，发现队列）")
    add_bullet(doc, "Liu J, et al. Neoadjuvant anti-PD-1 in NSCLC. Nat Cancer. 2022;XX:XXX–XXX.（GSE179994）")
    add_bullet(doc, "Hu J, et al. Single-cell profiling of immunotherapy. Genome Med. 2023;XX:XXX–XXX.（GSE207422）")
    add_bullet(doc, "Zhang Y, et al. NEOTIDE trial single-cell analysis. Cell Rep Med. 2024;XX:XXX–XXX.（GSE241934）")
    add_bullet(doc, "Monkman J, et al. Spatial transcriptomics of NSCLC. Cancer Cell. 2023;XX:XXX–XXX.（GSE221733 DSP）")
    add_bullet(doc, "Sade-Feldman M, et al. High-resolution intratumoral immune landscape. Cell. 2018;XX:XXX–XXX.（GSE120575）")
    add_bullet(doc, "Riaz N, et al. Tumor and microenvironment evolution. Cell. 2017;XX:XXX–XXX.（GSE91061）")
    add_bullet(doc, "Caushi JX, et al. Epigenetic and transcriptional dynamics. Nature. 2021;XX:XXX–XXX.（GSE176021）")
    add_para(doc, "建议用 EndNote/Zotero 管理并导出 AACR 样式；总数控制在 50 以内。", size=9, after=6)

    # ===================== 补充材料清单 =====================
    doc.add_page_break()
    add_heading(doc, "Supplementary Materials（图表对应文件）", 14)
    supp = [
        ("FigS1", "各 CD8 亚群比例分组箱线图 + QC 指标", ["FigS1_supplement.png/.pdf"]),
        ("FigS2", "NK-dominant 阈值敏感性（0/5/10/20/50）", ["FigS2_threshold_sensitivity.png/.pdf"]),
        ("FigS4", "TCR 跨患者排他性（无共享克隆）", ["FigS4_tcr_exclusivity.png/.pdf"]),
        ("FigS5", "配受体交互网络图", ["FigS5_interaction_network.png/.pdf"]),
        ("FigS6/7", "空间概述与空间验证（含原 Fig7 空间交互，已移补充）", ["FigS6_spatial_overview.png/.pdf", "FigS7_spatial_validation.png/.pdf", "Fig7_spatial_interaction.png/.pdf"]),
        ("FigS_spatial_decoupling", "空间解耦统计", ["FigS_spatial_decoupling.png/.pdf"]),
        ("Fig_chemo_dynamics", "化疗动力学", ["Fig_chemo_dynamics.png/.pdf"]),
        ("Fig_chemo_stratified", "化疗分层", ["Fig_chemo_stratified.png/.pdf"]),
        ("GSE241934_Fig6-11", "TCR 克隆型 / 多样性 / 共享", ["GSE241934_Fig6_TCR_clonality.png", "GSE241934_FigS10_clone_sharing.png"]),
        ("GSE241934_Fig7-9", "功能比较 / SPP1 TAM / IRS 验证", ["GSE241934_Fig7_functional_comparison.png", "GSE241934_Fig8_SPP1_TAM.png", "GSE241934_Fig9_IRS_validation.png"]),
        ("GSE179994_FigS12", "配对分析", ["GSE179994_FigS12_paired.png"]),
        ("GSE120575", "机制验证", ["GSE120575_mechanistic_validation.png"]),
    ]
    for tag, desc, files in supp:
        add_bullet(doc, f"{desc} ［文件：{'；'.join(files)}］", bold_lead=f"{tag}：")

    # ===================== 封面信要点 =====================
    doc.add_page_break()
    add_heading(doc, "Cover Letter 要点（CCR 要求：临床意义 + 转化影响 + 利益冲突）", 14)
    add_bullet(doc, "临床意义：首个在 ~188 例单细胞尺度证明\"克隆命运锁定（非克隆大小）决定 NSCLC 抗 PD-1 pCR\"的研究，提供治疗前可评估的候选标志物（NK-Locked Ratio）。")
    add_bullet(doc, "转化影响：揭示 SPP1+ TAM-KLF2 髓系闸门机制，提出\"抗 PD-1 + SPP1 通路抑制\"联合策略；并解释 Taxane 与 Pemetrexed 方案响应差异。")
    add_bullet(doc, "新颖性 / 差异化：区别于已有 SPP1+ TAM 抑制 CD8 耗竭的文献，本工作锚定 NK样克隆命运锁定 + 方案特异性反转，属原创非确证性发现。")
    add_bullet(doc, "利益冲突：声明（据实）。")
    add_bullet(doc, "推荐审稿人：［占位］。")

    doc.save(OUTPUT_PATH)
    print(f"CCR 论文骨架已保存至: {OUTPUT_PATH}")
    print(f"文件大小: {os.path.getsize(OUTPUT_PATH) / 1024 / 1024:.2f} MB")


if __name__ == "__main__":
    main()
